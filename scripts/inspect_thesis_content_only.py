from pathlib import Path
import sys

from docx import Document


DOC_PATH = Path("毕业论文") / "天津商业大学毕业设计(论文)张正豪v4.1-内容重做版.docx"
KEY_TERMS = [
    "时间银行",
    "闲置",
    "论坛",
    "导览",
    "校园",
    "跑腿",
    "地点热度",
    "小程序首页",
    "消息",
    "二手",
    "互助帖",
    "登录",
    "注册",
    "社区",
    "志愿者",
    "重点关怀",
    "异常预警",
    "便民信息",
    "数据库",
    "数据表",
    "Spring Boot",
    "Vue",
]


def print_hits(doc: Document) -> None:
    print("=== TERM HITS ===")
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        hits = [term for term in KEY_TERMS if term in text]
        if hits:
            safe_print(f"{idx:03d} {hits} :: {text}")


def print_ranges(doc: Document, ranges: list[tuple[int, int]]) -> None:
    print("=== RANGES ===")
    for start, end in ranges:
        print(f"--- {start}-{end} ---")
        for idx in range(start, min(end, len(doc.paragraphs))):
            para = doc.paragraphs[idx]
            style = getattr(para.style, "name", "")
            safe_print(f"{idx:03d} [{style}] {para.text}")


def print_tables(doc: Document, start: int = 0, end: int | None = None) -> None:
    print("=== TABLES ===")
    stop = len(doc.tables) if end is None else min(end, len(doc.tables))
    for table_idx in range(start, stop):
        table = doc.tables[table_idx]
        print(f"--- table {table_idx} rows={len(table.rows)} cols={len(table.columns)} ---")
        for row_idx, row in enumerate(table.rows[:8]):
            values = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            safe_print(f"{row_idx:02d} | {' | '.join(values)}")


def safe_print(text: str) -> None:
    sys.stdout.buffer.write((text + "\n").encode("utf-8", errors="replace"))


def main() -> None:
    doc = Document(str(DOC_PATH))
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    print_hits(doc)
    print_ranges(
        doc,
        [
            (60, 95),
            (108, 125),
            (140, 215),
            (230, 330),
            (535, 565),
        ],
    )
    print_tables(doc, 4, 11)


if __name__ == "__main__":
    main()

from pathlib import Path
from docx import Document


DOC_PATHS = [
    Path("毕业论文") / "天津商业大学毕业设计(论文)张正豪v4.1.docx",
    Path("毕业论文") / "天津商业大学毕业设计(论文)张正豪v4.1-5.1已修订.docx",
]


PARA_UPDATES = {
    44: "2.3 Vue 3与Element Plus\t7",
    85: "本课题旨在设计并实现一套基于 Spring Boot 3 + Vue 3 架构的社区公益服务对接管理平台，本系统的核心价值在于尝试建立一种可实现、可落地的“互联网+社区治理”新模式。",
    94: "本课题旨在响应国家“十五五”规划中关于健全基层社会治理体系的号召，针对当前社区公益服务中存在的供需错配、监管缺失和数据孤岛等痛点，结合社区治理与公益服务对接场景，基于 Spring Boot 3 和 Vue 3 设计并开发一套实时、规范、便捷的社区公益服务对接管理平台。具体目标包括：",
    134: "Spring Boot 3是由Pivotal团队提供的基于Spring框架的开源项目，旨在简化Spring应用的初始搭建和开发过程。Spring Boot 3全面支持Java 17及更高版本，并引入了诸多性能优化与新特性。通过Spring Security模块，可方便地实现RBAC权限控制、JWT无状态认证、密码加密等安全机制，非常适合本系统多角色权限管理的需求。",
    135: "2.3  Vue 3与Element Plus",
    154: "后端技术方面，Spring Boot 3作为Java生态中主流的单体/微服务应用开发框架，拥有完善的自动配置机制和快速开发能力。MyBatis-Plus作为MyBatis的增强工具，提供了开箱即用的 CRUD 接口和代码生成能力，能够显著降低数据访问层的开发复杂度。MySQL 8.0具备优秀的事务处理能力、查询性能和数据安全保障，完全能够满足本系统业务数据的存储与管理需求。Spring Security框架为RBAC权限控制和JWT认证提供了成熟方案，技术实现路径清晰。",
    155: "前端技术方面，Vue 3具有响应式开发和组件化组织优势。Element Plus是适配Vue 3的桌面端组件库，提供丰富、美观、易用的后台管理组件，能够快速构建高质量的管理界面。ECharts是数据可视化库，支持多种图表类型和灵活配置，能满足本系统数据看板的展示需求。",
    156: "在开发工具方面，VS Code拥有完善的插件生态，通过安装 Java、Spring Boot、Database 等扩展，可实现前后端项目的统一开发、调试和数据库管理，开发效率有保障。",
    145: "本系统的总体结构依据居民用户、志愿者用户、社区管理员和系统管理员的实际业务需求展开规划，整体可划分为移动端用户模块、管理端治理模块、后端业务服务模块和平台支撑模块四个部分。社区公益服务对接管理平台总体规划结构功能如图1所示。",
    148: "社区公益服务对接管理平台的目标群体主要包括需要发布求助的社区居民、通过认证后参与服务的志愿者，以及负责审核和治理的社区管理员。平台围绕“居民发布需求—管理员审核—志愿者认领—服务完成—评价反馈”的主流程展开，同时补充社区绑定、重点关怀对象管理、异常预警监测、AI 辅助草稿和数据可视化分析等能力，形成较完整的社区公益服务闭环。",
    258: "根据 MySQL 关系型数据库的设计原则，并结合 `schema_v2_prd.sql` 中的实际建表结果，当前系统数据库共包含 27 张表，覆盖用户、社区、公益服务、预警治理和系统支撑等多个方面。考虑论文展示篇幅，下文选取 7 张最能体现当前业务主线的核心表结构进行说明。",
}


TABLE_DATA = {
    4: [
        ["字段名称", "标识符", "数据类型", "说明"],
        ["id", "√", "bigint unsigned", "主键ID"],
        ["username", "-", "varchar(50)", "登录用户名"],
        ["password_hash", "-", "varchar(100)", "密码哈希"],
        ["role", "-", "tinyint unsigned", "角色：超级管理员/社区管理员/普通用户"],
        ["community_id", "-", "bigint unsigned", "所属社区ID"],
        ["community_join_status", "-", "tinyint unsigned", "社区加入状态"],
        ["real_name", "-", "varchar(50)", "真实姓名"],
        ["phone", "-", "varchar(20)", "联系电话"],
        ["email", "-", "varchar(100)", "邮箱"],
        ["avatar_url", "-", "varchar(255)", "头像地址"],
        ["gender", "-", "tinyint unsigned", "性别"],
        ["address", "-", "varchar(255)", "常住地址"],
        ["status", "-", "tinyint unsigned", "账号状态：0禁用 1启用"],
        ["points", "-", "bigint", "志愿积分"],
        ["last_login_at", "-", "datetime(3)", "最近登录时间"],
        ["created_at", "-", "datetime(3)", "创建时间"],
        ["updated_at", "-", "datetime(3)", "更新时间"],
    ],
    5: [
        ["字段名称", "标识符", "数据类型", "说明"],
        ["id", "√", "bigint unsigned", "主键ID"],
        ["user_id", "-", "bigint unsigned", "申请用户ID"],
        ["community_id", "-", "bigint unsigned", "申请加入社区ID"],
        ["invite_code", "-", "varchar(32)", "使用的邀请码"],
        ["real_name", "-", "varchar(50)", "申请人姓名"],
        ["phone", "-", "varchar(20)", "联系电话"],
        ["address", "-", "varchar(255)", "居住地址"],
        ["status", "-", "tinyint unsigned", "状态：待审核/通过/拒绝/撤回"],
        ["reviewer_user_id", "-", "bigint unsigned", "审核人ID"],
        ["reviewed_at", "-", "datetime(3)", "审核时间"],
        ["reject_reason", "-", "varchar(255)", "拒绝原因"],
        ["created_at", "-", "datetime(3)", "创建时间"],
        ["updated_at", "-", "datetime(3)", "更新时间"],
        ["", "", "", ""],
    ],
    6: [
        ["字段名称", "标识符", "数据类型", "说明"],
        ["id", "√", "bigint unsigned", "主键ID"],
        ["user_id", "-", "bigint unsigned", "用户ID"],
        ["community_id", "-", "bigint unsigned", "所属社区ID"],
        ["cert_status", "-", "tinyint unsigned", "认证状态"],
        ["skill_tags", "-", "json", "技能标签集合"],
        ["service_radius_km", "-", "decimal(6,2)", "服务半径"],
        ["available_time", "-", "varchar(255)", "可服务时间说明"],
        ["avg_rating", "-", "decimal(4,2)", "平均评分"],
    ],
    7: [
        ["字段名称", "标识符", "数据类型", "说明"],
        ["id", "√", "bigint unsigned", "主键ID"],
        ["user_id", "-", "bigint unsigned", "用户ID"],
        ["community_id", "-", "bigint unsigned", "所属社区ID"],
        ["care_type", "-", "varchar(64)", "关怀类型"],
        ["care_level", "-", "tinyint unsigned", "关怀等级"],
        ["living_status", "-", "varchar(64)", "居住情况"],
        ["health_note", "-", "varchar(500)", "健康备注"],
        ["emergency_contact_name", "-", "varchar(64)", "紧急联系人姓名"],
        ["emergency_contact_phone", "-", "varchar(32)", "紧急联系人电话"],
        ["monitor_enabled", "-", "tinyint unsigned", "是否启用异常监测"],
    ],
    8: [
        ["字段名称", "标识符", "数据类型", "说明"],
        ["id", "√", "bigint unsigned", "主键ID"],
        ["requester_user_id", "-", "bigint unsigned", "需求发布人ID"],
        ["community_id", "-", "bigint unsigned", "所属社区ID"],
        ["service_type", "-", "varchar(50)", "服务类型"],
        ["description", "-", "text", "需求描述"],
        ["service_address", "-", "varchar(255)", "服务地址"],
        ["urgency_level", "-", "tinyint unsigned", "紧急程度"],
        ["status", "-", "tinyint unsigned", "状态：待审核至已完成"],
        ["audit_by_user_id", "-", "bigint unsigned", "审核人ID"],
    ],
    9: [
        ["字段名称", "标识符", "数据类型", "说明"],
        ["id", "√", "bigint unsigned", "主键ID"],
        ["request_id", "-", "bigint unsigned", "需求ID"],
        ["volunteer_user_id", "-", "bigint unsigned", "志愿者用户ID"],
        ["claim_status", "-", "tinyint unsigned", "认领状态"],
        ["claim_at", "-", "datetime(3)", "认领时间"],
    ],
    10: [
        ["字段名称", "标识符", "数据类型", "说明"],
        ["id", "√", "bigint unsigned", "主键ID"],
        ["claim_id", "-", "bigint unsigned", "认领记录ID"],
        ["request_id", "-", "bigint unsigned", "需求ID"],
        ["resident_user_id", "-", "bigint unsigned", "居民用户ID"],
        ["volunteer_user_id", "-", "bigint unsigned", "志愿者用户ID"],
        ["evaluator_role", "-", "tinyint unsigned", "评价方角色"],
        ["rating", "-", "tinyint unsigned", "评分：1-5"],
        ["content", "-", "varchar(500)", "评价内容"],
    ],
}


def fill_table(table, rows):
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value


def main() -> None:
    updated = []
    skipped = []
    for path in DOC_PATHS:
        if not path.exists():
            continue
        try:
            doc = Document(str(path))
            for idx, text in PARA_UPDATES.items():
                doc.paragraphs[idx].text = text

            for table_idx, rows in TABLE_DATA.items():
                fill_table(doc.tables[table_idx], rows)

            doc.save(str(path))
            updated.append(str(path))
        except PermissionError:
            skipped.append(str(path))

    print("UPDATED:")
    for item in updated:
        print(item)
    print("SKIPPED:")
    for item in skipped:
        print(item)


if __name__ == "__main__":
    main()

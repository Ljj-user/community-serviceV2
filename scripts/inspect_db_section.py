from pathlib import Path
from docx import Document


DOC = Path("毕业论文") / "天津商业大学毕业设计(论文)张正豪v4.1.docx"


def main() -> None:
    doc = Document(str(DOC))
    for i in range(248, 340):
        para = doc.paragraphs[i]
        print(f"{i:03d} [{getattr(para.style, 'name', '')}] {para.text}")

    print(f"TABLES={len(doc.tables)}")
    for idx, table in enumerate(doc.tables[:20]):
        rows = len(table.rows)
        cols = len(table.columns)
        first = " | ".join(cell.text for cell in table.rows[0].cells) if rows else ""
        print(f"T{idx}: {rows}x{cols} :: {first}")
        if idx >= 4:
            for ridx, row in enumerate(table.rows[:6]):
                print("   ", ridx, " | ".join(cell.text for cell in row.cells))


if __name__ == "__main__":
    main()

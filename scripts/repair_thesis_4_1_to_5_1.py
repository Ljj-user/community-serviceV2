from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document


DOC_NAME = "天津商业大学毕业设计(论文)张正豪v4.1-5.1已修订.docx"
SRC_DOC_NAME = "天津商业大学毕业设计(论文)张正豪v4.1.docx"


TEXT_UPDATES = {
    68: "关键词：社区公益服务；志愿者管理；SpringBoot；Vue3；数字化治理；社区治理",
    90: "相比之下，国外的社区数字化探索更多源于其成熟的社区自治文化，主要体现为社区连接强化、志愿协同机制完善以及服务治理数字化水平较高等特征。以美国的 Nextdoor 为代表的邻里社区平台，通过严格的住址验证机制构建高信任度线上社区；而 VolunteerMatch 等志愿服务平台则更强调志愿者招募、任务匹配与服务连接效率。这些实践说明，数字平台若能兼顾真实身份、社区关系与服务流程，就能够更有效地促进居民参与和社区互助。",
    91: "综上所述，国外在社区连接、志愿协同和数字治理方面提供了宝贵经验，但若直接照搬，容易因国情、文化和基层治理机制差异而水土不服；而国内现有系统又往往偏重行政管理，缺乏对民生微需求的持续响应能力。因此，开发一套贴合本土社区治理场景、兼顾服务对接效率与后台治理能力的社区公益服务平台，既具有现实必要性，也具有一定的研究价值。",
    116: "导言。主要介绍社区公益服务对接管理平台的设计与实现的背景与意义，阐述国内外社区数字治理与公益服务平台的发展现状，并提出本文的研究内容与方法。",
    308: "移动端首页承担平台导览与快速分流作用，页面由顶部社区栏、轮播区、公告提示、快捷功能卡片和 AI 输入入口组成。用户进入首页后，可先查看当前绑定社区、浏览公告，再根据需要进入“我要发布求助”“我去帮助别人”“志愿者认证”或“便民信息”等功能。",
    199: "系统结构模块设计是在需求分析基础上，对平台的功能边界、角色职责和数据流转关系进行拆分与归纳，为后续编码实现提供结构化蓝图。本系统采用前后端分离的 B/S 架构，整体由移动端、管理端、后端服务层与 MySQL 数据层组成，并通过统一 RESTful API 实现数据交互。",
    201: "本系统整体架构围绕“居民发布求助、管理员审核、志愿者接单、服务完成、评价反馈、异常预警”的主业务闭环展开。移动端主要承担居民与志愿者的日常操作入口，管理端主要承担社区治理、审核与统计分析职责，后端统一负责权限校验、业务编排、数据持久化和审计留痕，其整体结构如图6所示。",
    205: "在整体架构基础上，平台主要功能模块可分为用户接入、社区绑定、公益需求、志愿者认证、服务评价、公告与便民信息、重点关怀对象、异常预警、数据看板和系统治理等部分，如图7所示。",
    209: "表现层通过 Axios 调用后端 RESTful API，并以 JSON 作为主要数据交换格式。业务逻辑层基于 Spring Boot、Spring Security 和 Service 组件实现权限控制、状态流转和审核处理，数据访问层基于 MyBatis-Plus 完成表对象映射与查询封装，数据层由 MySQL 统一存储用户、社区、需求、认领、评价和治理数据。该分层设计使界面、业务与数据职责清晰，也便于后续模块扩展与维护。",
    211: "用户端功能模块包括：登录注册与社区加入、移动端首页、服务大厅、需求发布、我的需求、我的服务、服务评价、志愿者认证、社区公告、便民信息、AI 助手和个人中心。其中，移动端采用统一账户体系，用户既可以发布求助，也可以在通过志愿者认证后参与接单服务。",
    214: "在用户端功能模块划分基础上，其功能活动关系主要体现为：用户完成注册登录后先绑定所属社区，再进入首页浏览公告与快捷入口；居民可发布求助并跟踪处理进度；志愿者在通过认证后可进入服务大厅接单、完成服务并沉淀评价信息，其活动关系如图9所示。",
    218: "社区管理员端功能模块包括：数据看板、社区加入审核、需求审核、志愿者审核、重点关怀对象管理、异常预警处理、公告管理、便民信息管理、用户管理和审计日志查看。社区管理员主要聚焦本社区数据范围内的审核、治理与服务运行监控。",
    222: "在社区管理员端功能模块划分基础上，其功能活动关系表现为：管理员先审核居民加入社区申请，再审核公益需求和志愿者认证申请，同时通过公告发布、重点关怀对象维护和异常预警处置，对社区公益服务运行过程进行持续治理，如图11所示。",
    228: "系统管理员端功能模块包括：全局数据看板、系统配置、社区与用户全局治理、审计日志、数据导出与备份等，同时兼具社区管理员全部业务权限。与社区管理员相比，系统管理员更强调跨社区统计、平台级配置和整体运行监管。",
    232: "在系统管理员端功能模块划分基础上，其功能活动关系体现为：系统管理员可从平台级视角查看多社区运行数据，统筹维护系统参数、审计关键操作、处理全局治理事项，并在必要时接管社区管理员功能，其活动关系如图13所示。",
    239: "公益需求状态流转模块是本系统最核心的业务流程之一，用于支撑社区公益服务从“需求产生”到“服务完成”再到“评价沉淀”的全流程管理。该模块既面向居民与志愿者提供服务协同入口，也为社区管理员提供审核与过程监管抓手。",
    240: "具体流程如下：居民在完成登录和社区绑定后提交需求，需求首先进入“待审核”状态；社区管理员审核通过后，需求进入“待接单”或“可认领”状态；通过认证的志愿者可在服务大厅查看并认领任务，任务随即进入“已认领”或“服务中”状态；志愿者完成服务后提交完成信息，居民确认后进入“待评价”阶段；评价完成后，服务记录、认领记录与评价数据统一沉淀到数据库，并可进一步参与统计分析和异常预警判断。",
    241: "该流程实现了居民、志愿者与管理员三方在同一业务链路中的协同，也保证了每条公益服务需求都具备可审核、可追踪、可评价、可复盘的状态记录，构成系统平台的核心闭环。",
    245: "4.2.2  社区加入与后台审核治理流程设计",
    246: "社区加入与后台审核治理流程用于保证平台的数据边界与社区归属关系正确。用户注册后默认并不直接拥有社区数据访问权限，而是需要通过邀请码校验或提交加入申请的方式与具体社区建立绑定关系；管理员审核通过后，系统将用户的 community_id 写入账户信息，后续公告、便民信息、公益需求和治理数据均按该字段进行隔离。",
    247: "在治理流程上，社区加入审核、需求审核和志愿者审核共同构成后台人工把关机制。也就是说，系统并不把关键节点完全交给自动流转，而是在加入社区、发布需求和获得志愿者资格等环节设置人工确认步骤，从而降低错误绑定、虚假求助和无效接单等风险，保证平台运行秩序。",
    248: "4.3  数据库存储设计",
    249: "系统开发当中，数据库存储设计是一个非常重要的环节，它直接影响到项目的性能、稳定性和可扩展性。本系统采用 MySQL 8 作为关系型数据库，后端通过 MyBatis-Plus 完成数据持久化与对象映射，本小节将结合当前业务主线，对数据库概念模型与核心表结构进行说明。",
    251: "数据库概念模型设计用于抽象系统中的核心实体及其关系。结合当前社区公益服务业务场景，平台的主要实体包括用户、社区区域、社区加入申请、志愿者认证档案、重点关怀对象、公益服务需求、服务认领记录、服务评价、社区公告、便民信息、异常预警事件和审计日志等。各实体围绕 community_id、user_id、request_id、claim_id 等关键字段建立关联，共同支撑前后端功能实现。",
    252: "从业务角度看，系统数据库表可归纳为用户与权限、社区绑定与治理、公益服务流程、公共信息发布和后台运行支撑五类。其中，用户既是需求发布者，也可以在认证后成为志愿者；管理员通过审核、公告、预警和审计等表结构完成社区治理。根据不同角色与功能需求分析，抽象得到的平台 E-R 关系如图15所示。",
    258: "根据 MySQL 关系型数据库的设计原则，并结合 `schema_v2_prd.sql` 中的实际建表结果，当前系统数据库共包含 27 张表，覆盖用户、社区、公益服务、预警治理和系统支撑等多个方面。本小节重点说明与论文主题最相关的核心表结构。",
    259: "用户表 `sys_user`：用于存储统一账户信息、角色、认证状态、所属社区、联系方式和账户状态等，是移动端和管理端权限控制的基础。",
    260: "社区绑定相关表：`community_join_application` 用于记录用户加入社区的申请、审核状态和处理结果；`community_invite_code` 用于生成社区邀请码；`sys_region` 用于维护区域与社区层级结构，共同支撑社区绑定与数据隔离。",
    261: "志愿者认证相关表：`volunteer_profile` 用于保存志愿者技能标签、服务半径、可服务时间和认证状态，支撑志愿者审核、展示与接单能力；相关技能信息可通过辅助表继续沉淀。",
    262: "重点关怀对象表 `care_subject_profile`：用于存储独居老人、残障居民等重点关怀对象的档案信息、关怀等级、紧急联系人和监测开关等内容，为社区主动关怀与异常预警提供基础数据。",
    263: "公益需求表 `service_request`：用于保存居民发布的公益求助信息，包括服务类型、紧急程度、地址、描述、审核状态和所属社区，是业务主表之一。",
    264: "服务认领表 `service_claim`：用于记录志愿者认领、执行和完成服务的过程信息，并与具体需求建立关联，是需求流转到服务落实的关键中间表。",
    265: "服务评价表 `service_evaluation`：用于保存居民对服务结果的评分与文字反馈，并通过 `claim_id` 与认领记录建立一对一对应关系，形成服务闭环的结果沉淀。",
    266: "公共信息相关表：`announcement` 用于发布社区公告，`convenience_info` 用于维护药店、医院、服务热线等便民信息，共同支撑社区日常信息触达。",
    267: "治理与审计相关表：`anomaly_alert_event` 用于记录超时未认领、超时未完成等异常事件，`audit_log` 用于记录后台关键操作日志，支撑系统追溯、治理留痕与风险分析。",
    268: "除上述核心表外，系统还通过 `community_banner`、`service_request_tag`、`volunteer_match_snapshot`、`sys_notification`、`sys_config` 和 `backup_record` 等表补充首页展示、标签匹配、消息通知与系统运维能力。",
    269: "上述数据库设计既满足了“发布、审核、接单、完成、评价”的主流程，也兼顾了“社区绑定、重点关怀、预警治理和后台审计”的扩展需求，为系统实现提供了稳定的数据基础。",
    270: "对于论文展示而言，`sys_user`、`community_join_application`、`volunteer_profile`、`care_subject_profile`、`service_request`、`service_claim`、`service_evaluation`、`announcement`、`convenience_info`、`anomaly_alert_event` 和 `audit_log` 等核心表最能体现系统当前的业务主线。",
    271: "上述数据库设计既满足了“发布、审核、接单、完成、评价”的主流程，也兼顾了“社区绑定、重点关怀、预警治理和后台审计”的扩展需求，为系统实现提供了稳定的数据基础。",
}


HEADING_2 = {198, 237, 248, 274, 544, 550}
HEADING_3 = {200, 210, 217, 227, 238, 245, 250, 257, 276, 450}
HEADING_4 = {277, 322, 345, 362, 394, 420, 432, 451, 471, 490, 516}
BODY_LEADERS = {
    278, 290, 323, 329, 346, 351, 363, 367, 395, 400, 421, 427,
    433, 438, 453, 459, 472, 478, 491, 497, 517, 523,
}
BODY_FORCE = {
    325, 331, 364, 397, 423, 428, 435, 455, 460, 474, 493, 519,
}


def resolve_doc_path() -> Path:
    thesis_dir = Path.cwd() / "毕业论文"
    exact = thesis_dir / DOC_NAME
    if exact.exists():
        return exact
    matches = sorted(thesis_dir.glob("*v4.1-5.1*已修订*.docx"))
    if not matches:
        raise FileNotFoundError("未找到目标论文文档。")
    return matches[0]


def resolve_src_doc_path() -> Path:
    thesis_dir = Path.cwd() / "毕业论文"
    exact = thesis_dir / SRC_DOC_NAME
    if exact.exists():
        return exact
    matches = sorted(thesis_dir.glob("*v4.1.docx"))
    if not matches:
        raise FileNotFoundError("未找到论文基准文档。")
    return matches[0]


def copy_paragraph_format(src, dst) -> None:
    src_fmt = src.paragraph_format
    dst_fmt = dst.paragraph_format
    dst.style = src.style
    dst.alignment = src.alignment
    dst_fmt.left_indent = src_fmt.left_indent
    dst_fmt.right_indent = src_fmt.right_indent
    dst_fmt.first_line_indent = src_fmt.first_line_indent
    dst_fmt.keep_together = src_fmt.keep_together
    dst_fmt.keep_with_next = src_fmt.keep_with_next
    dst_fmt.page_break_before = src_fmt.page_break_before
    dst_fmt.widow_control = src_fmt.widow_control
    dst_fmt.space_before = src_fmt.space_before
    dst_fmt.space_after = src_fmt.space_after
    dst_fmt.line_spacing = src_fmt.line_spacing
    dst_fmt.line_spacing_rule = src_fmt.line_spacing_rule


def has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//*[local-name()='drawing']"))


def set_text(paragraph, text: str) -> None:
    if paragraph.text != text:
        paragraph.text = text


def rebuild_tail_from_source(target_doc: Document, source_doc: Document) -> None:
    target_anchor = next(
        p for p in target_doc.paragraphs if p.text.strip() == "5  结论"
    )
    source_anchor = next(
        p for p in source_doc.paragraphs if p.text.strip() == "5  结论"
    )

    target_body = target_doc._body._body
    source_body = source_doc._body._body

    target_children = list(target_body)
    source_children = list(source_body)

    target_start = target_children.index(target_anchor._element)
    source_start = source_children.index(source_anchor._element)

    for child in target_children[target_start:]:
        target_body.remove(child)

    for child in source_children[source_start:]:
        target_body.append(deepcopy(child))


def update_conclusion_sections(doc: Document) -> None:
    paragraphs = doc.paragraphs
    heading_2_template = next(p for p in paragraphs if p.text.strip() == "4.1  系统结构模块设计")
    body_template = next(p for p in paragraphs if p.text.startswith("系统结构模块设计是在需求分析基础上"))

    def find_index(text: str) -> int:
        for idx, p in enumerate(paragraphs):
            if p.text.strip() == text:
                return idx
        raise ValueError(f"未找到段落: {text}")

    h51 = find_index("5.1  本文所做的主要工作")
    h52 = find_index("5.2  今后进一步研究的方向")
    ref = find_index("参考文献")

    copy_paragraph_format(heading_2_template, paragraphs[h51])
    copy_paragraph_format(heading_2_template, paragraphs[h52])

    p51 = [
        "本文围绕“社区公益服务对接管理平台”的建设目标，完成了从需求分析、系统设计到前后端实现与功能联调的完整过程。系统以“居民发布求助—管理员审核—志愿者接单—服务完成—居民评价”为主线，同时补充了社区绑定、重点关怀、异常预警和后台治理能力，形成了较完整的社区公益服务闭环。本小节归纳总结本文所做工作主要为以下几个方面：",
        "完成了系统总体方案设计与核心业务建模，围绕社区公益服务场景，梳理并设计了登录注册、社区加入、需求发布、需求审核、志愿者认证、服务认领、完成确认、评价反馈等核心业务流程，并明确了基于 community_id 的社区归属与数据隔离规则；",
        "完成了移动端与管理端的主要功能实现。移动端实现了登录/注册与社区加入、首页、服务大厅、需求发布、我的需求、我的服务、评价反馈、志愿者认证、社区公告、便民信息和 AI 助手等页面；管理端实现了数据看板、社区加入审核、需求审核、志愿者审核、重点关怀对象管理、异常预警、用户管理和审计日志等页面；",
        "完成了关键技术落地与系统联调。前端基于 Vue 3、Pinia 和组件化设计实现页面复用与状态协同，后端基于 Spring Boot、Spring Security 和 MyBatis-Plus 提供权限控制、业务接口与数据持久化能力，支撑前后端一体化运行；",
        "形成了贴合社区治理场景的功能亮点。系统在基础服务对接之外，引入了社区绑定审核、重点关怀对象建档、异常预警监测和 AI 草稿辅助等能力，增强了平台在社区公益服务对接、治理留痕和风险处置方面的实用价值。",
    ]

    p52 = [
        "社区公益服务对接管理平台的基本功能已基本完成，能够满足社区辖内求助发布、审核治理、志愿服务对接和评价反馈等核心业务需求。但从真实社区运行场景来看，平台在智能化辅助、精细化治理和持续运营能力方面仍有进一步优化空间，后续可从以下几个方向继续完善：",
        "（1）功能深化方面：在现有求助发布、服务认领和评价反馈基础上，可继续完善需求标签推荐、志愿者能力匹配、服务进度提醒和异常超时预警联动机制，提高供需匹配效率与服务闭环质量。",
        "（2）社区治理方面：可进一步强化社区加入审核、重点关怀对象动态维护和多角色数据看板能力，并结合真实社区反馈补充更细的治理指标，使管理员能够更及时地发现高风险对象、积压需求和服务瓶颈。",
        "（3）智能辅助方面：可继续优化 AI 草稿辅助和分析记录能力，使其更好地服务于需求整理、公告编写和后台研判等场景，但仍应坚持“AI 辅助、人工审核把关”的原则，避免将关键治理决策完全交由自动化模块处理。",
    ]

    for offset, text in enumerate(p51, start=1):
        set_text(paragraphs[h51 + offset], text)
        copy_paragraph_format(body_template, paragraphs[h51 + offset])

    for offset, text in enumerate(p52, start=1):
        idx = h52 + offset
        if idx >= ref:
            break
        set_text(paragraphs[idx], text)
        copy_paragraph_format(body_template, paragraphs[idx])


def normalize_section_format(doc: Document) -> None:
    paragraphs = doc.paragraphs
    title_1_template = paragraphs[196]
    heading_2_template = paragraphs[198]
    heading_3_template = paragraphs[200]
    heading_4_template = paragraphs[277]
    body_template = paragraphs[199]
    caption_template = paragraphs[204]

    copy_paragraph_format(title_1_template, paragraphs[196])

    for idx in HEADING_2:
        copy_paragraph_format(heading_2_template, paragraphs[idx])

    for idx in HEADING_3:
        copy_paragraph_format(heading_3_template, paragraphs[idx])

    for idx in HEADING_4:
        copy_paragraph_format(heading_4_template, paragraphs[idx])

    for idx in range(199, 272):
        paragraph = paragraphs[idx]
        text = paragraph.text.strip()
        if not text or idx in HEADING_2 or idx in HEADING_3 or idx in HEADING_4:
            continue
        if text.startswith("图 "):
            copy_paragraph_format(caption_template, paragraph)
        else:
            copy_paragraph_format(body_template, paragraph)

    for idx in range(274, 555):
        paragraph = paragraphs[idx]
        text = paragraph.text.strip()
        if idx in HEADING_2 or idx in HEADING_3 or idx in HEADING_4:
            continue
        if not text:
            continue
        if text.startswith("图 "):
            copy_paragraph_format(caption_template, paragraph)
            continue
        if has_drawing(paragraph):
            continue
        if (
            idx in BODY_FORCE
            or idx in BODY_LEADERS
            or text.startswith("（1）")
            or text.startswith("（2）")
            or text.startswith("（3）")
            or text.startswith("关键代码如下")
        ):
            copy_paragraph_format(body_template, paragraph)

    for idx in range(545, 555):
        if paragraphs[idx].text.strip():
            copy_paragraph_format(body_template, paragraphs[idx])

    if len(paragraphs) > 555:
        paragraphs[555].alignment = 1


def main() -> None:
    doc_path = resolve_doc_path()
    src_doc_path = resolve_src_doc_path()

    doc = Document(str(doc_path))
    src_doc = Document(str(src_doc_path))

    rebuild_tail_from_source(doc, src_doc)

    for idx, text in TEXT_UPDATES.items():
        set_text(doc.paragraphs[idx], text)

    update_conclusion_sections(doc)
    normalize_section_format(doc)
    doc.save(str(doc_path))
    print(doc_path)


if __name__ == "__main__":
    main()

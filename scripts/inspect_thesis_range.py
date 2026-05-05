from pathlib import Path
import sys

from docx import Document


DOC_PATH = Path("毕业论文") / "天津商业大学毕业设计(论文)张正豪v4.1-内容重做版.docx"


def safe_print(text: str) -> None:
    sys.stdout.buffer.write((text + "\n").encode("utf-8", errors="replace"))


def main() -> None:
    doc = Document(str(DOC_PATH))
    start = 320
    end = 540
    for idx in range(start, min(end, len(doc.paragraphs))):
        para = doc.paragraphs[idx]
        style = getattr(para.style, "name", "")
        safe_print(f"{idx:03d} [{style}] {para.text}")


if __name__ == "__main__":
    main()

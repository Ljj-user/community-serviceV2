from pathlib import Path
from docx import Document


DOC = Path("毕业论文") / "天津商业大学毕业设计(论文)张正豪v4.1.docx"
RANGES = [
    (0, 90),
    (108, 120),
    (140, 210),
    (540, 620),
]


def main() -> None:
    doc = Document(str(DOC))
    for start, end in RANGES:
        print(f"RANGE {start}-{end}")
        for idx in range(start, min(end, len(doc.paragraphs))):
            p = doc.paragraphs[idx]
            style = getattr(p.style, "name", "")
            print(f"{idx:03d} [{style}] {p.text}")


if __name__ == "__main__":
    main()
